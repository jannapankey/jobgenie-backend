from docx import Document
import os

def save_resume_as_docx(resume_json, filename):
    document = Document()

    # Extract fields
    full_name = resume_json.get('full_name', 'Unknown Name')
    email = resume_json.get('email', 'unknown@example.com')
    phone = resume_json.get('phone', 'Unknown Phone')
    summary = resume_json.get('summary', '')
    education = resume_json.get('education', '')
    skills = resume_json.get('skills', [])
    experience_list = resume_json.get('experience', [])

    # Add Full Name at top
    document.add_heading(full_name, level=0)
    document.add_paragraph(f"Email: {email} | Phone: {phone}")

    # Add a line break
    document.add_paragraph()

    # Professional Summary
    document.add_heading('Professional Summary', level=1)
    document.add_paragraph(summary)

    # Education
    document.add_heading('Education', level=1)
    document.add_paragraph(education)

    # Skills
    document.add_heading('Skills', level=1)
    if isinstance(skills, list):
        skills_text = ', '.join(skills)
    else:
        skills_text = str(skills)
    document.add_paragraph(skills_text)

    # Experience
    document.add_heading('Experience', level=1)
    for exp in experience_list:
        title = exp.get('title', 'No Title')
        company = exp.get('company', 'No Company')
        dates = exp.get('dates', 'No Dates')
        bullets = exp.get('bullets', [])

        document.add_heading(f"{title} at {company} ({dates})", level=2)
        for bullet in bullets:
            document.add_paragraph(bullet, style='List Bullet')

    # Ensure downloads folder exists
    downloads_dir = os.path.join(os.getcwd(), "downloads")
    os.makedirs(downloads_dir, exist_ok=True)

    # Save document
    safe_filename = filename.replace(" ", "_")
    file_path = os.path.join(downloads_dir, safe_filename)
    document.save(file_path)

    return file_path
