# app.py

from flask import Flask, request, jsonify, render_template, send_from_directory
import os
import json
from agent import run_resume_agent
from docx import Document
import mammoth
import uuid

app = Flask(__name__)

# API route: Generate resume
@app.route("/generate", methods=["POST"])
def generate_resume():
    try:
        data = request.get_json()

        full_name = data.get("full_name")
        email = data.get("email")
        phone = data.get("phone")
        education = data.get("education")
        skills = data.get("skills")
        work_experiences = data.get("work_experiences", [])
        job_description = data.get("job_description")

        candidate_info = {
            "full_name": full_name,
            "email": email,
            "phone": phone,
            "education": education,
            "skills": skills,
            "work_experiences": work_experiences
        }

        # Run full multi-agent process
        final_resume_json = run_resume_agent(candidate_info, job_description)

        # Create .docx file
        doc = Document()

        doc.add_heading(full_name, 0)
        doc.add_paragraph(f"{email} | {phone}")
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(final_resume_json["summary"])

        doc.add_heading("Skills", level=1)
        for skill in final_resume_json["skills"]:
            doc.add_paragraph(skill, style="List Bullet")

        doc.add_heading("Education", level=1)
        doc.add_paragraph(final_resume_json["education"])

        doc.add_heading("Experience", level=1)
        for job in final_resume_json["experience"]:
            doc.add_paragraph(f"{job['title']} at {job['company']} ({job['dates']})", style="Heading 2")
            for bullet in job["bullets"]:
                doc.add_paragraph(bullet, style="List Bullet")

        # Save document
        unique_id = str(uuid.uuid4())
        filename = f"{unique_id}.docx"
        filepath = os.path.join("downloads", filename)
        os.makedirs("downloads", exist_ok=True)
        doc.save(filepath)

        # Convert .docx to HTML for preview
        with open(filepath, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            resume_html = result.value

        download_link = f"/download/{filename}"

        return jsonify({
            "resume_text": resume_html,
            "download_link": download_link
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

# API route: Serve resume downloads
@app.route("/download/<path:filename>")
def download_resume(filename):
    return send_from_directory("downloads", filename, as_attachment=True)

# Web UI
@app.route("/")
def index():
    return render_template("index.html")

if __name__ == "__main__":
    app.run(debug=True)
